# -*- coding: utf-8 -*-
"""
《基于 Spring Boot 与 Vue 的多合一环境监测保护系统的设计与实现》
Word 文档生成脚本（国内高校通用排版规范）：
- 页面：A4，上2.5cm 下2.5cm 左3.0cm 右2.5cm
- 正文：宋体小四(12pt)，西文 Times New Roman，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号(16pt) 居中；二级标题：黑体四号(14pt)；三级标题：黑体小四(12pt)
- 表题/图题：黑体五号(10.5pt) 居中
- 表格内容：宋体五号(10.5pt)
- 页眉：论文题目（宋体五号居中）；页脚：页码
- 封面（用户可填写占位）、中英摘要、目录域、正文、参考文献
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import paper_content as PC

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "多合一环境监测保护系统设计与实现_课程设计论文.docx")


def set_run_font(run, cn_font, size_pt, bold=False, ascii_font="Times New Roman"):
    """设置 run 的中英文字体、字号、加粗。"""
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), cn_font)


def add_para(doc, text, cn_font="宋体", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_indent=0, line_spacing=1.5, space_before=0, space_after=0, ascii_font="Times New Roman",
             outline_level=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    if line_spacing:
        pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent:
        pf.first_line_indent = Pt(first_indent)
    if outline_level is not None:
        pPr = p._p.get_or_add_pPr()
        outlineLvl = OxmlElement('w:outlineLvl')
        outlineLvl.set(qn('w:val'), str(outline_level))
        pPr.append(outlineLvl)
    run = p.add_run(text)
    set_run_font(run, cn_font, size, bold, ascii_font)
    return p


def set_cell(cell, text, cn_font="宋体", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, ascii_font="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = align
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, cn_font, size, bold, ascii_font)
    # 垂直居中
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def add_table(doc, data):
    """data: {title, header, rows}"""
    if data.get("title"):
        add_para(doc, data["title"], cn_font="黑体", size=10.5, bold=False,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
    header = data["header"]
    rows = data["rows"]
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(header):
        set_cell(table.rows[0].cells[j], h, cn_font="黑体", size=10.5, bold=False)
    # 内容
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell(table.rows[i + 1].cells[j], val, cn_font="宋体", size=10.5)
    # 表后空一行
    add_para(doc, "", size=6)


def add_code(doc, text, size=9):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.1
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.6)
        run = p.add_run(line if line else " ")
        set_run_font(run, "宋体", size, False, "Consolas")


def add_fig_placeholder(doc, text):
    add_para(doc, text, cn_font="黑体", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=6, space_after=10)


def set_header_footer(section, header_text=None, page_start=None, page_num_style=None):
    """配置页眉页脚。header_text 为空则无页眉。"""
    if header_text:
        section.header.is_linked_to_previous = False
        hp = section.header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(header_text)
        set_run_font(run, "宋体", 10.5)
        # 页眉下边框
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = ""

    # 页脚页码
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if page_num_style:
        fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE \\* MERGEFORMAT '
        fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
        run = fp.add_run()
        run._r.append(fld_begin)
        run2 = fp.add_run()
        run2._r.append(instr)
        run3 = fp.add_run()
        run3._r.append(fld_end)
        for r in (run, run2, run3):
            set_run_font(r, "宋体", 10.5)

    if page_start is not None:
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        pgNumType.set(qn('w:start'), str(page_start))


def add_toc_field(doc):
    """插入目录域（Word 中打开后右键更新域即可生成目录）。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（目录：请在 Word 中右键此处选择“更新域”生成）"
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run = p.add_run()
    run._r.append(fld_begin)
    run2 = p.add_run(); run2._r.append(instr)
    run3 = p.add_run(); run3._r.append(fld_sep)
    run4 = p.add_run(); run4._r.append(t)
    run5 = p.add_run(); run5._r.append(fld_end)
    for r in (run, run2, run3, run4, run5):
        set_run_font(r, "宋体", 12)


def add_keywords(doc, text, label, size=12, en=False):
    """关键词行：标签黑体加粗，内容宋体。en=True 时全部用 Times New Roman。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)
    pf.space_before = Pt(6)
    label_run = p.add_run(label)
    set_run_font(label_run, "Times New Roman" if en else "黑体", size, bold=True)
    content_run = p.add_run(text)
    set_run_font(content_run, "Times New Roman" if en else "宋体", size, bold=False)


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()

    # ===== 页面设置 =====
    section0 = doc.sections[0]
    section0.page_height = Cm(29.7)
    section0.page_width = Cm(21.0)
    section0.top_margin = Cm(2.5)
    section0.bottom_margin = Cm(2.5)
    section0.left_margin = Cm(3.0)
    section0.right_margin = Cm(2.5)

    # 默认样式：正文 宋体小四
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ===== 封面（第 1 节，无页眉页脚） =====
    c = PC.COVER
    # 顶部留白
    add_para(doc, "", size=24); add_para(doc, "", size=24)
    add_para(doc, c["school"], cn_font="黑体", size=26, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_para(doc, c["doc_type"], cn_font="黑体", size=22, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_before=12)
    add_para(doc, "", size=20)
    # 题目
    add_para(doc, PC.PAPER_TITLE, cn_font="黑体", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_before=10, space_after=10)
    add_para(doc, "", size=20)
    # 信息行
    for label, val in [("课　程", c["course"]), ("学　院", c["school"]),
                       ("专　业", c["major_class"]), ("班　级", c["major_class"]),
                       ("学　号", c["student_id"]), ("姓　名", c["student"]),
                       ("指导教师", c["teacher"]), ("完成日期", c["date"])]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        r1 = p.add_run(label + "：")
        set_run_font(r1, "黑体", 14, bold=True)
        r2 = p.add_run(val)
        set_run_font(r2, "宋体", 14)
    add_page_break(doc)

    # ===== 第 2 节：摘要 + Abstract + 目录 =====
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec2.top_margin = Cm(2.5); sec2.bottom_margin = Cm(2.5)
    sec2.left_margin = Cm(3.0); sec2.right_margin = Cm(2.5)
    set_header_footer(sec2, header_text=None, page_start=1, page_num_style=True)

    # 中文摘要
    add_para(doc, PC.PAPER_TITLE, cn_font="黑体", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=10)
    add_para(doc, "摘　要", cn_font="黑体", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=6)
    for para in PC.ABSTRACT_CN:
        add_para(doc, para, cn_font="宋体", size=12, first_indent=24)
    add_keywords(doc, PC.ABSTRACT_CN_KEYWORDS.split("：", 1)[1], "关键词：")
    add_page_break(doc)

    # 英文摘要
    add_para(doc, PC.PAPER_TITLE_EN, cn_font="Times New Roman", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=10, ascii_font="Times New Roman")
    add_para(doc, "Abstract", cn_font="Times New Roman", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=6)
    for para in PC.ABSTRACT_EN:
        add_para(doc, para, cn_font="Times New Roman", size=12, first_indent=24)
    add_keywords(doc, PC.ABSTRACT_EN_KEYWORDS.split(":", 1)[1].strip(), "Keywords: ", en=True)
    add_page_break(doc)

    # 目录
    add_para(doc, "目　录", cn_font="黑体", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=10)
    add_toc_field(doc)
    add_page_break(doc)

    # ===== 第 3 节：正文（页眉 + 页码重排从 1） =====
    sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec3.top_margin = Cm(2.5); sec3.bottom_margin = Cm(2.5)
    sec3.left_margin = Cm(3.0); sec3.right_margin = Cm(2.5)
    set_header_footer(sec3, header_text=PC.PAPER_TITLE, page_start=1, page_num_style=True)

    for item in PC.CONTENT:
        kind = item[0]
        if kind == "h1":
            add_para(doc, item[1], cn_font="黑体", size=16, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                     space_before=12, space_after=10, outline_level=0)
        elif kind == "h2":
            add_para(doc, item[1], cn_font="黑体", size=14, bold=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                     space_before=8, space_after=6, outline_level=1)
        elif kind == "h3":
            add_para(doc, item[1], cn_font="黑体", size=12, bold=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                     space_before=6, space_after=4, outline_level=2)
        elif kind == "body":
            add_para(doc, item[1], cn_font="宋体", size=12, first_indent=24)
        elif kind == "caption":
            add_para(doc, item[1], cn_font="黑体", size=10.5,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)
        elif kind == "table":
            add_table(doc, item[1])
        elif kind == "code":
            add_code(doc, item[1])
        elif kind == "fig":
            add_fig_placeholder(doc, item[1])
        else:
            raise ValueError("未知元素类型: " + kind)

    # ===== 参考文献 =====
    add_para(doc, "参考文献", cn_font="黑体", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
             space_before=12, space_after=10)
    for i, ref in enumerate(PC.REFERENCES, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(2)
        pf.left_indent = Cm(0.8)
        pf.first_line_indent = Cm(-0.8)  # 悬挂缩进
        run = p.add_run("[%d] %s" % (i, ref))
        set_run_font(run, "宋体", 10.5)

    doc.save(OUT)
    # 设置打开时自动更新域（目录自动生成）
    settings = doc.settings.element
    updateFields = OxmlElement('w:updateFields')
    updateFields.set(qn('w:val'), 'true')
    settings.append(updateFields)
    doc.save(OUT)
    print("已生成:", OUT)


if __name__ == "__main__":
    build()
