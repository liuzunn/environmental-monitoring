# -*- coding: utf-8 -*-
"""深度校验 docx 排版：字体/字号/缩进/行距/页眉页脚/节设置。"""
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

d = docx.Document('多合一环境监测保护系统设计与实现_课程设计论文.docx')

def run_font(run):
    rPr = run._element.find(qn('w:rPr'))
    ea = ascii_f = None
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            ea = rf.get(qn('w:eastAsia'))
            ascii_f = rf.get(qn('w:ascii'))
    return ea, ascii_f

print('===== 节设置 =====')
for i, s in enumerate(d.sections):
    print('节%d: 页宽%.1fcm 高%.1fcm 上%.1f 下%.1f 左%.1f 右%.1f' % (
        i + 1, s.page_width.cm, s.page_height.cm, s.top_margin.cm,
        s.bottom_margin.cm, s.left_margin.cm, s.right_margin.cm))
    has_hdr = s.header.is_linked_to_previous
    hdr_text = s.header.paragraphs[0].text if s.header.paragraphs else ''
    print('   页眉(is_linked=%s): "%s"' % (has_hdr, hdr_text[:30]))

print('===== 关键段落字体抽样 =====')
for p in d.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    pf = p.paragraph_format
    r = p.runs[0] if p.runs else None
    ea, af = run_font(r) if r is not None else (None, None)
    sz = r.font.size.pt if (r and r.font.size) else None
    fi = pf.first_line_indent.pt if pf.first_line_indent else 0
    ls = pf.line_spacing
    if t.startswith('第 1 章'):
        print('一级标题 | %s | 字体:%s/%s 字号:%s 行距:%s' % (t, ea, af, sz, ls))
    elif t.startswith('1.1'):
        print('二级标题 | %s | 字体:%s/%s 字号:%s' % (t, ea, af, sz))
    elif t.startswith('4.3.1'):
        print('三级标题 | %s | 字体:%s/%s 字号:%s' % (t, ea, af, sz))
    elif t.startswith('随着工业化'):
        print('正文段 | %s... | 字体:%s/%s 字号:%s 首行缩进:%.0fpt 行距:%s' % (t[:12], ea, af, sz, fi, ls))
        break

print('===== 页脚页码检查 =====')
for i, s in enumerate(d.sections):
    flds = s.footer.paragraphs[0]._p.xml.count('PAGE')
    print('节%d 页脚含PAGE域次数: %d' % (i + 1, flds))

print('===== 表格数 =====', len(d.tables))
print('===== 关键词加粗检查 =====')
for p in d.paragraphs:
    if p.text.startswith('关键词：'):
        for r in p.runs:
            print('  run: "%s" 黑体=%s 加粗=%s' % (r.text[:12], run_font(r)[0], r.font.bold))
        break
